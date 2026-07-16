"""
Document converter.

Turns every supported file in `app/resources/` into Markdown stored in
`app/markdown/` with the same base name (per the user guide).

Supported extensions: .docx, .pdf, .txt, .md, .html, .htm

Idempotent — re-runs skip files whose .md output is newer than the source.
"""
from pathlib import Path

from docx import Document
import fitz                                 # pymupdf
from bs4 import BeautifulSoup
import html2text

from app.utils.logger import log_info, log_success, log_warning, log_error


RESOURCES_DIR = Path("app/resources")
MARKDOWN_DIR = Path("app/markdown")

MARKDOWN_DIR.mkdir(parents=True, exist_ok=True)


# =========================================================
# Per-type converters
# =========================================================

def _convert_docx(file_path: Path) -> str:
    """DOCX → markdown, preserving headings + table content."""
    doc = Document(file_path)
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            parts.append("")
            continue

        style = (paragraph.style.name or "").lower()
        if "heading 1" in style:
            parts.append(f"# {text}")
        elif "heading 2" in style:
            parts.append(f"## {text}")
        elif "heading 3" in style:
            parts.append(f"### {text}")
        else:
            parts.append(text)

    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def _convert_pdf(file_path: Path) -> str:
    pdf = fitz.open(file_path)
    pages: list[str] = []
    for index, page in enumerate(pdf, start=1):
        pages.append(f"\n## Page {index}\n")
        pages.append(page.get_text())
    pdf.close()
    return "\n".join(pages)


def _convert_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _convert_md_passthrough(file_path: Path) -> str:
    """Markdown files copy through unchanged."""
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _convert_html(file_path: Path) -> str:
    """HTML → markdown via html2text after BeautifulSoup cleanup."""
    raw = file_path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    converter = html2text.HTML2Text()
    converter.body_width = 0
    converter.ignore_links = False
    converter.ignore_images = True
    return converter.handle(str(soup))


# =========================================================
# Dispatcher
# =========================================================

_DISPATCH = {
    ".docx": _convert_docx,
    ".pdf":  _convert_pdf,
    ".txt":  _convert_txt,
    ".md":   _convert_md_passthrough,
    ".html": _convert_html,
    ".htm":  _convert_html,
}


def _convert_single(file_path: Path) -> Path | None:
    suffix = file_path.suffix.lower()
    if suffix not in _DISPATCH:
        log_warning(f"Skipping unsupported file: {file_path.name}")
        return None

    output = MARKDOWN_DIR / f"{file_path.stem}.md"

    # Idempotent skip
    if output.exists() and output.stat().st_mtime >= file_path.stat().st_mtime:
        log_info(f"Up to date: {file_path.name}")
        return output

    try:
        content = _DISPATCH[suffix](file_path)
    except Exception as e:
        log_error(f"Convert {file_path.name}", e)
        return None

    output.write_text(content, encoding="utf-8")
    log_success(f"Converted: {file_path.name} -> {output.name}")
    return output


def process_resources() -> list[Path]:
    if not RESOURCES_DIR.exists():
        log_warning(f"Resources directory not found: {RESOURCES_DIR}")
        return []

    outputs: list[Path] = []
    for file_path in sorted(RESOURCES_DIR.iterdir()):
        if file_path.is_dir():
            continue
        output = _convert_single(file_path)
        if output is not None:
            outputs.append(output)

    log_success(f"Conversion complete — {len(outputs)} file(s) ready")
    return outputs
